"""
Claude API-integration för Alejandro
"""
import os
import json
import base64
import re
import anthropic
from loguru import logger
from prompts.system import SYSTEM_PROMPT, bygg_user_prompt

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

def analysera_reflektion_for_lankar(text: str) -> dict:
    """
    Analysera reflektionstexten för länkar och filtyper som Alejandro inte kan hantera.
    Returnerar dict med flagga och orsak.
    """
    import re
    if not text:
        return {"flagga": False, "orsak": None}

    flagga_monster = [
        (r'youtube\.com|youtu\.be',                    "YouTube-video"),
        (r'vimeo\.com',                                  "Vimeo-video"),
        (r'tiktok\.com',                                 "TikTok-video"),
        (r'instagram\.com',                              "Instagram-inlägg"),
        (r'facebook\.com',                               "Facebook-inlägg"),
        (r'docs\.google\.com',                          "Google Docs/Slides/Sheets"),
        (r'drive\.google\.com',                         "Google Drive-länk"),
        (r'onedrive\.live\.com|1drv\.ms',              "OneDrive-länk"),
        (r'dropbox\.com',                                "Dropbox-länk"),
        (r'\.mp4|\.mov|\.avi|\.mkv|\.webm',          "videofil"),
        (r'\.mp3|\.wav|\.m4a|\.ogg|\.aac',          "ljudfil"),
        (r'\.pptx|\.ppt',                               "PowerPoint-presentation"),
        (r'\.xlsx|\.xls',                               "Excel-fil"),
        (r'\.zip|\.rar|\.7z',                          "komprimerad fil"),
    ]

    hittade = []
    for monster, beskrivning in flagga_monster:
        if re.search(monster, text, re.IGNORECASE):
            hittade.append(beskrivning)

    if hittade:
        orsak = "Inlämningen innehåller material som kräver mänsklig granskning: " + ", ".join(hittade)
        return {"flagga": True, "orsak": orsak}
    return {"flagga": False, "orsak": None}

def analysera_filtyp_kan_ej_hanteras(filtyp: str, filnamn: str = "") -> dict:
    """Kolla om en bifogad fil är av en typ Alejandro inte kan hantera"""
    EJ_HANTERBARA = [
        ("video/", "videofil"),
        ("audio/", "ljudfil"),
        ("application/vnd.ms-powerpoint", "PowerPoint"),
        ("application/vnd.openxmlformats-officedocument.presentationml", "PowerPoint"),
        ("application/vnd.ms-excel", "Excel"),
        ("application/vnd.openxmlformats-officedocument.spreadsheetml", "Excel"),
        ("application/zip", "ZIP-arkiv"),
        ("application/x-rar", "RAR-arkiv"),
    ]
    
    EJ_HANTERBARA_EXT = [
        (".mp4", "videofil"), (".mov", "videofil"), (".avi", "videofil"),
        (".mp3", "ljudfil"), (".wav", "ljudfil"), (".m4a", "ljudfil"),
        (".pptx", "PowerPoint"), (".ppt", "PowerPoint"),
        (".xlsx", "Excel"), (".xls", "Excel"),
        (".zip", "ZIP-arkiv"), (".rar", "RAR-arkiv"),
    ]
    
    for prefix, namn in EJ_HANTERBARA:
        if filtyp.lower().startswith(prefix):
            return {"flagga": True, "orsak": f"Bifogad {namn} kräver mänsklig granskning"}
    
    for ext, namn in EJ_HANTERBARA_EXT:
        if filnamn.lower().endswith(ext):
            return {"flagga": True, "orsak": f"Bifogad {namn} kräver mänsklig granskning"}
    
    return {"flagga": False, "orsak": None}


def docx_till_pdf_base64(data_bytes: bytes) -> str | None:
    """Konvertera Word-dokument till PDF och returnera som base64"""
    try:
        import io
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT

        # Läs Word-dokumentet
        doc = Document(io.BytesIO(data_bytes))
        
        # Extrahera text
        texter = []
        for para in doc.paragraphs:
            if para.text.strip():
                texter.append(para.text)

        # Skapa PDF i minnet
        pdf_buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=A4,
            leftMargin=20*mm,
            rightMargin=20*mm,
            topMargin=20*mm,
            bottomMargin=20*mm,
        )
        
        styles = getSampleStyleSheet()
        story = []
        for text in texter:
            story.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["Normal"]))
            story.append(Spacer(1, 4*mm))
        
        if not story:
            return None
            
        pdf_doc.build(story)
        pdf_bytes = pdf_buffer.getvalue()
        
        import base64 as b64lib
        result = b64lib.standard_b64encode(pdf_bytes).decode("utf-8")
        logger.info(f"✅ Word → PDF konverterat ({len(pdf_bytes)//1024}KB)")
        return result
        
    except Exception as e:
        logger.warning(f"Kunde inte konvertera Word till PDF: {e}")
        return None



_jwt_token_cache = None

def get_jwt_token() -> str | None:
    """Logga in som Alejandro och hämta JWT-token"""
    global _jwt_token_cache
    if _jwt_token_cache:
        return _jwt_token_cache
    import requests
    backend = os.getenv("BACKEND_URL", "https://eportfolj-hemundervisning-production.up.railway.app")
    email = os.getenv("AI_HANDLAGGARE_EMAIL", "alejandro@hemundervisning.ax")
    losenord = os.getenv("AI_LOSENORD", "")
    try:
        resp = requests.post(
            f"{backend}/api/auth/logga-in",
            json={"email": email, "losenord": losenord},
            timeout=10
        )
        if resp.status_code == 200:
            _jwt_token_cache = resp.json().get("token")
            logger.info("✅ Alejandro inloggad via API")
            return _jwt_token_cache
        else:
            logger.warning(f"JWT-inloggning misslyckades ({resp.status_code})")
    except Exception as e:
        logger.debug(f"JWT-inloggning fel: {e}")
    return None

def ladda_fil_via_inlamning(barn_id: int, moment_id: int, inlamning_id: int) -> tuple[bytes, str] | None:
    """Ladda fil via /api/trad/barn/{barnId}/moment/{momentId}/fil/{inlamningId}"""
    import requests
    backend = os.getenv("BACKEND_URL", "https://eportfolj-hemundervisning-production.up.railway.app")
    token = get_jwt_token()
    headers = {"Authorization": f"Bearer {token}"} if token else {}
    url = f"{backend}/api/trad/barn/{barn_id}/moment/{moment_id}/fil/{inlamning_id}"
    logger.debug(f"Laddar fil via trad: {url}")
    try:
        resp = requests.get(url, headers=headers, timeout=15)
        if resp.status_code == 200:
            content_type = resp.headers.get("content-type", "application/octet-stream").split(";")[0].strip()
            logger.info(f"✅ Fil laddad via trad-endpoint ({len(resp.content)//1024}KB)")
            return resp.content, content_type
        logger.warning(f"Kunde inte ladda via trad ({resp.status_code}): {url}")
    except Exception as e:
        logger.warning(f"Fel vid trad-filnedladdning: {e}")
    return None

def ladda_fil_fran_server(lagrad_path: str) -> tuple[bytes, str] | None:
    """
    Ladda ner en fil från backenden via autentiserat API-anrop.
    lagrad_path är t.ex. /app/uploads/inlamningar/abc123.jpg
    """
    import requests
    import tempfile
    import mimetypes

    backend = os.getenv("BACKEND_URL", "https://eportfolj-hemundervisning-production.up.railway.app")
    
    # Konvertera lokal serverväg till API-endpoint
    # /app/uploads/inlamningar/fil.jpg -> /api/filer/inlamningar/fil.jpg
    if lagrad_path.startswith("/app/uploads/"):
        api_path = lagrad_path.replace("/app/uploads/", "/api/filer/")
    elif lagrad_path.startswith("/app/"):
        api_path = lagrad_path.replace("/app/", "/api/filer/")
    else:
        api_path = f"/api/filer/{lagrad_path.lstrip('/')}"

    token = get_jwt_token()
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"

    try:
        url = f"{backend}{api_path}"
        logger.debug(f"Laddar fil: {url}")
        resp = requests.get(url, headers=headers, timeout=15)
        if resp.status_code == 200:
            content_type = resp.headers.get("content-type", "application/octet-stream")
            media_type = content_type.split(";")[0].strip()
            return resp.content, media_type
        else:
            logger.warning(f"Kunde inte ladda fil ({resp.status_code}): {url}")
    except Exception as e:
        logger.warning(f"Fel vid filnedladdning: {e}")
    return None

def ladda_bild_som_base64(fil_url: str) -> tuple[str, str] | None:
    """Försök ladda en bild och konvertera till base64"""
    BILDTYPER = ["image/jpeg", "image/png", "image/gif", "image/webp"]
    
    # Lokal serverväg - ladda via autentiserat API
    if not fil_url.startswith("http"):
        result = ladda_fil_fran_server(fil_url)
        if result:
            data, media_type = result
            if media_type in BILDTYPER:
                return base64.standard_b64encode(data).decode("utf-8"), media_type
        return None
    
    # Publik URL
    try:
        import requests
        resp = requests.get(fil_url, timeout=10)
        if resp.status_code == 200:
            media_type = resp.headers.get("content-type", "image/jpeg").split(";")[0].strip()
            if media_type not in BILDTYPER:
                return None
            return base64.standard_b64encode(resp.content).decode("utf-8"), media_type
    except Exception as e:
        logger.warning(f"Kunde inte ladda bild: {e}")
    return None

def bedom_inlamning(
    inlamning: dict,
    barn: dict,
    familj: dict,
    historik: list,
    laroplan: list,
) -> dict | None:
    """
    Skicka inlämning till Claude och få tillbaka en bedömning.
    Returnerar parsed JSON-svar eller None vid fel.
    """
    user_prompt = bygg_user_prompt(inlamning, barn, familj, historik, laroplan)
    
    messages = []
    fil_ej_laddad = False
    
    # Bygg meddelande – med eller utan bild
    fil_url = inlamning.get("fil_url")
    fil_typ = inlamning.get("fil_typ", "") or ""

    content_blocks = []

    if fil_url:
        # Prova trad-endpoint först om vi har barn_id och moment_id
        fil_data = None
        barn_id_for_fil = inlamning.get("_barn_id_for_fil")
        moment_id_for_fil = inlamning.get("_moment_id_for_fil")
        inlamning_id_for_fil = inlamning.get("id")
        
        if barn_id_for_fil and moment_id_for_fil and inlamning_id_for_fil:
            fil_data = ladda_fil_via_inlamning(barn_id_for_fil, moment_id_for_fil, inlamning_id_for_fil)
        
        # Fallback: ladda via lagrad_path
        if fil_data is None:
            fil_data = ladda_fil_fran_server(fil_url) if not fil_url.startswith("http") else None
        if fil_data is None and fil_url.startswith("http"):
            try:
                import requests as req
                r = req.get(fil_url, timeout=15)
                if r.status_code == 200:
                    fil_data = (r.content, r.headers.get("content-type", "").split(";")[0].strip())
            except Exception:
                pass

        if not fil_data:
            logger.warning(f"Kunde inte ladda fil – bedömer utan bilaga: {fil_url}")
            content_blocks.append({
                "type": "text",
                "text": f"[En fil är bifogad ({inlamning.get('fil_typ','okänd typ')}) men kunde inte laddas för granskning. Notera detta i bedömningen och be familjen bekräfta att filen är korrekt uppladdad om det behövs.]"
            })
        if fil_data:
            data_bytes, media_type = fil_data
            b64 = base64.standard_b64encode(data_bytes).decode("utf-8")

            if media_type in ["image/jpeg", "image/png", "image/gif", "image/webp"]:
                logger.info(f"📷 Skickar bild ({media_type})")
                content_blocks.append({
                    "type": "image",
                    "source": {"type": "base64", "media_type": media_type, "data": b64}
                })

            elif media_type == "application/pdf":
                logger.info("📄 Skickar PDF")
                content_blocks.append({
                    "type": "document",
                    "source": {"type": "base64", "media_type": "application/pdf", "data": b64}
                })

            elif media_type in ["text/plain", "text/csv"] or fil_url.endswith(".txt"):
                try:
                    text_content = data_bytes.decode("utf-8", errors="replace")
                    logger.info(f"📝 Skickar textfil ({len(text_content)} tecken)")
                    content_blocks.append({
                        "type": "text",
                        "text": f"[Bifogad textfil]\n{text_content[:8000]}"
                    })
                except Exception:
                    logger.warning("Kunde inte läsa textfil")

            elif media_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ] or fil_url.endswith((".doc", ".docx")):
                logger.info("📝 Konverterar Word → PDF")
                pdf_b64 = docx_till_pdf_base64(data_bytes)
                if pdf_b64:
                    content_blocks.append({
                        "type": "document",
                        "source": {"type": "base64", "media_type": "application/pdf", "data": pdf_b64}
                    })
                else:
                    # Fallback: extrahera ren text
                    try:
                        import io
                        from docx import Document
                        doc = Document(io.BytesIO(data_bytes))
                        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                        content_blocks.append({"type": "text", "text": f"[Word-dokument]\n{text[:8000]}"})
                        logger.info("📝 Word-text extraherad som fallback")
                    except Exception:
                        content_blocks.append({"type": "text", "text": "[Word-dokument bifogat – kunde ej konverteras]"})
            else:
                logger.info(f"⏭ Okänd filtyp: {media_type}")
                content_blocks.append({
                    "type": "text",
                    "text": f"[Bifogad fil av typen {media_type} – kan ej visas direkt, men familjen har bifogat en fil som bevis]"
                })
        else:
            logger.warning(f"Bifogad fil kunde inte laddas – flaggar för mänsklig granskning")
            return None, True  # Flagga till mänsklig handläggare

    content_blocks.append({"type": "text", "text": user_prompt})

    if len(content_blocks) > 1:
        messages.append({"role": "user", "content": content_blocks})
    else:
        messages.append({"role": "user", "content": user_prompt})

    try:
        response = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=1500,
            system=SYSTEM_PROMPT,
            messages=messages,
        )
        
        raw = response.content[0].text.strip()
        logger.debug(f"Claude råsvar: {raw[:500]}")
        
        # Extrahera JSON (även om Claude råkat lägga text runt)
        json_match = re.search(r'\{.*\}', raw, re.DOTALL)
        if json_match:
            parsed = json.loads(json_match.group())
            return parsed
        else:
            logger.error(f"Kunde inte hitta JSON i svaret: {raw[:200]}")
            return None
            
    except json.JSONDecodeError as e:
        logger.error(f"JSON-parsningsfel: {e}")
        return None
    except anthropic.APIError as e:
        logger.error(f"Anthropic API-fel: {e}")
        return None
    except Exception as e:
        logger.error(f"Oväntat fel i bedom_inlamning: {e}")
        return None
